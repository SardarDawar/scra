import streamlit as st
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches

def extract_content(url):
    # Make a request to the webpage
    response = requests.get(url)

    # Parse the HTML content of the webpage with BeautifulSoup
    soup = BeautifulSoup(response.content, 'html.parser')

    # Find the footer element by its HTML tag or class
    footer = soup.find('footer')
    navbar = soup.find('nav', {'class': 'navbar'})

    # Remove the footer element and all its contents from the parsed HTML
    if footer:
        footer.decompose()

    if navbar:
        navbar.decompose()

    # Extract the remaining content of the webpage
    content = soup.get_text()

    # Remove blank lines
    lines = [line for line in content.splitlines() if line.strip()]

    # Join the remaining lines
    text = '\n'.join(lines)

    title = soup.title.string

    # Create a new Word document
    document = Document()

    # Add a title to the document
    document.add_heading(f"{title}", 0)

    # Add a paragraph to the document with the extracted content
    document.add_paragraph(text)

    # Save the document to a file
    document.save('content.docx')

    return title

st.title("Webpage Content Extractor")
url = st.text_input("Enter the URL of the webpage you want to extract content from:")

if st.button("Extract Content"):
    if url:
        title = extract_content(url)
        st.write(f"{title} content saved in content.docx file")
    else:
        st.write("Please enter a valid URL.")
